# -*- coding: utf-8 -*-
"""Build an updated DOCX report from the current outputs/ data."""

from __future__ import annotations

import csv
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "智能媒体工作坊_技术报告模板..docx"
OUT = ROOT / "基础案例A-视觉分析2" / "技术报告_选题4_5_视觉模型拓展.docx"
FALLBACK_OUT = ROOT / "基础案例A-视觉分析2" / "技术报告_选题4_5_视觉模型拓展_完善版.docx"

TOPIC4 = ROOT / "outputs" / "topic4"
TOPIC5 = ROOT / "outputs" / "topic5"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING_COLOR = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "E8EEF5"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def run_font(run, size=10.5, bold=False, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def paragraph(doc, text: str, indent=True, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    run_font(r, size=size)
    return p


def heading(doc, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    run_font(r, size=14 if level == 1 else 12, bold=True, color=HEADING_COLOR)
    return p


def caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    run_font(r, size=9, color=MUTED)


def picture(doc, path: Path, cap: str, width=5.6):
    if not path.exists():
        paragraph(doc, f"（缺少图像：{path.name}）", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def delete_paragraph(p):
    element = p._element
    element.getparent().remove(element)
    p._p = p._element = None


def shade(cell, fill=TABLE_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def cell_text(cell, text, bold=False, size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_borders(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    pr = table._tbl.tblPr
    borders = pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C3D0")


def add_table(doc, headers: list[str], rows: list[list[object]], cap: str, first_left=True):
    caption(doc, cap)
    table = doc.add_table(rows=1, cols=len(headers))
    table_borders(table)
    for cell, text in zip(table.rows[0].cells, headers):
        cell_text(cell, text, bold=True, size=8.8)
        shade(cell)
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            align = WD_ALIGN_PARAGRAPH.LEFT if first_left and i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(cell, value, size=8.4, align=align)
    return table


def fill_cover(doc):
    ps = doc.paragraphs
    if len(ps) < 14:
        return
    ps[3].text = "从端正到扭曲：不同书写目标下的符号变形规律与 AI 识别局限"
    ps[8].text = "本报告基于当前 outputs 数据完善，合并完成选题4，并在此基础上拓展完成选题5："
    ps[12].text = "☑ ④ 不同书写目标（端正易读 / 美观 / 极致扭曲）下的视觉差异分析"
    ps[13].text = "☑ ⑤ 基于美观 vs 极致扭曲的多模型识别、形态特征解释与 AI 局限分析"
    for i in range(14):
        for run in ps[i].runs:
            run_font(run, size=16 if i in (0, 2, 3) else 11, bold=i in (0, 2, 3))


def fmt(x: str | float, digits=3) -> str:
    return f"{float(x):.{digits}f}"


def build():
    topic4_metrics = read_csv(TOPIC4 / "metrics" / "topic4_classification_metrics.csv")[0]
    topic4_dist = read_csv(TOPIC4 / "tables" / "topic4_group_distances.csv")
    topic5_metrics = read_csv(TOPIC5 / "tables" / "topic5_model_performance.csv")
    top_features = read_csv(TOPIC5 / "tables" / "topic5_fusion_feature_importance.csv")[:15]

    doc = Document(TEMPLATE)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
    fill_cover(doc)
    for p in list(doc.paragraphs)[14:]:
        delete_paragraph(p)

    doc.add_page_break()
    heading(doc, "摘要", 1)
    paragraph(
        doc,
        "本报告围绕《从端正到扭曲：不同书写目标下的符号变形规律与 AI 识别局限》展开。选题4首先基于端正易读、美观书写、极致扭曲三类裁切后的抄写格图像，复用基础案例A的 Data 分组、视觉特征提取、PCA 和 t-SNE 分析流程，使用本地 ResNet50 权重提取深度视觉特征，检验三类书写目标是否存在显著视觉差异。选题5在选题4的基础上进一步聚焦最容易混淆的“美观 vs 极致扭曲”，比较 ResNet50、Swin_T、Morphology 和 Fusion 四种特征方案，分析 AI 对细粒度符号变形的识别能力及局限。当前统一输出采用 page_id 分组划分训练集和测试集，降低同一扫描页样本同时进入训练与测试的风险。",
    )
    paragraph(
        doc,
        f"结果显示，选题4三分类 Accuracy 为 {fmt(topic4_metrics['accuracy'])}，Macro F1 为 {fmt(topic4_metrics['macro_f1'])}；三组距离中“美观 vs 极致扭曲”的中心距离最小，说明其是更困难的细粒度区分对象。选题5中 Fusion 融合特征 Accuracy 为 {fmt([r for r in topic5_metrics if r['model']=='Fusion'][0]['accuracy'])}，Macro F1 为 {fmt([r for r in topic5_metrics if r['model']=='Fusion'][0]['macro_f1'])}，优于单独深度特征和单独形态学特征。该结果表明，深度模型能够捕捉整体视觉风格，但解释扭曲机制仍需要外接框、轮廓复杂度、参考偏离等可解释形态特征补充。",
    )
    paragraph(doc, "关键词：符号变形；书写目标；ResNet50；Swin Transformer；形态学特征；AI 识别局限", indent=False)

    heading(doc, "1 数据与已有资源检查", 1)
    paragraph(
        doc,
        "本项目没有直接使用整张扫描页作为分析对象，而是先将三类扫描页裁切为单个抄写格图像。裁切后的目录为基础案例A的 Data 目录，三类子文件夹分别为端正易读、美观、极致扭曲，每组 2880 张，共 8640 张。裁切索引保存在 Data/image_index.csv，包含 image_path、label、source_page、page_id、crop_id、row、col 等字段。",
    )
    rows = [
        ["端正易读", "2880", "3-以端正易读方式抄写符号", "Data/端正易读"],
        ["美观", "2880", "4-以美观方式抄写符号", "Data/美观"],
        ["极致扭曲", "2880", "5-以极致扭曲方式抄写符号", "Data/极致扭曲"],
    ]
    add_table(doc, ["类别", "裁切图像数", "原始扫描页来源", "整理后目录"], rows, "表1  数据整理结果")
    paragraph(
        doc,
        "本地模型资源位于视觉特征计算/model。当前检测到 ResNet50.pth、Swin_T.pth、ViT_B16.pth，以及 CLIP 和 DINOv2 的 HuggingFace 本地缓存。正式实验优先使用 ResNet50 与 Swin_T 的 torchvision 本地权重；ViT_B16、CLIP 和 DINOv2 可作为后续扩展，但未作为本次主实验结论来源。",
    )

    heading(doc, "2 视觉分析.py 的复用与兼容处理", 1)
    paragraph(
        doc,
        "基础案例A中的视觉分析.py 已实现按 Data 子文件夹读取类别、生成 image_index.csv、提取视觉特征、PCA 降维、t-SNE/UMAP 二维嵌入和可视化等核心流程。该流程被本项目保留并复用。由于原脚本中仍包含 ViTFeatureExtractor 等较旧 transformers 接口，且当前任务强调离线运行，本报告采用兼容脚本使用本地 torchvision ResNet50 和 Swin_T 权重完成正式实验。原始视觉分析.py 已备份为视觉分析_backup.py，未直接覆盖。",
    )

    heading(doc, "3 选题4：三类书写目标的视觉差异", 1)
    paragraph(
        doc,
        "选题4的目标是证明端正易读、美观、极致扭曲三类书写方式在视觉特征空间中是否存在差异。实验使用 ResNet50 深度特征，经 PCA 降维后训练 RandomForestClassifier，并使用 t-SNE 观察二维分布。分类划分使用 GroupShuffleSplit 按 page_id 分组进行 8:2 划分，尽量避免同页样本泄漏。",
    )
    picture(doc, TOPIC4 / "figures" / "topic4_sample_examples.png", "图1  三类书写目标的裁切样例", 5.8)
    picture(doc, TOPIC4 / "figures" / "topic4_resnet50_pca_scatter.png", "图2  选题4 ResNet50 PCA 二维分布", 5.5)
    picture(doc, TOPIC4 / "figures" / "topic4_resnet50_tsne_scatter.png", "图3  选题4 ResNet50 t-SNE 二维分布", 5.5)
    dist_rows = [
        [
            r["pair"],
            fmt(r["centroid_distance_top10PC"]),
            fmt(r["centroid_distance_tSNE"]),
            fmt(r["PC1_cliffs_delta"]),
        ]
        for r in topic4_dist
    ]
    add_table(doc, ["对比", "前10 PC中心距离", "t-SNE中心距离", "PC1 Cliff's delta"], dist_rows, "表2  三类书写目标的视觉特征距离")
    add_table(
        doc,
        ["模型", "样本数", "训练数", "测试数", "Accuracy", "Macro F1"],
        [[
            "ResNet50 + PCA + RandomForest",
            topic4_metrics["n_samples"],
            topic4_metrics["n_train"],
            topic4_metrics["n_test"],
            fmt(topic4_metrics["accuracy"]),
            fmt(topic4_metrics["macro_f1"]),
        ]],
        "表3  选题4三分类性能",
    )
    picture(doc, TOPIC4 / "figures" / "topic4_confusion_matrix.png", "图4  选题4三分类混淆矩阵", 4.5)
    paragraph(
        doc,
        "从距离结果看，端正易读与极致扭曲之间的前10 PC中心距离最大，端正易读与美观次之，美观与极致扭曲最小。该结果说明三类书写目标确实存在系统性视觉差异，但美观与极致扭曲在特征空间中更接近，也为选题5聚焦二者提供了依据。",
    )

    heading(doc, "4 选题5：美观与极致扭曲的细粒度识别", 1)
    paragraph(
        doc,
        "选题5建立在选题4的结论之上，不再重复证明三类差异，而是选择最难区分的“美观 vs 极致扭曲”作为拓展对象。该任务比较四种特征方案：ResNet50 深度视觉特征、Swin_T 深度视觉特征、Morphology 可解释形态学特征，以及三者拼接后的 Fusion 融合特征。形态学特征包括墨迹占比、外接框宽度/高度/面积比例、轮廓周长、复杂度、边缘密度、方向熵和 ref_chamfer 等指标。",
    )
    picture(doc, TOPIC5 / "figures" / "topic5_aesthetic_vs_distorted_examples.png", "图5  美观与极致扭曲裁切样例对比", 5.8)
    perf_rows = [
        [
            r["model"],
            r["n_features"],
            r["n_samples"],
            r["n_test"],
            fmt(r["accuracy"]),
            fmt(r["macro_f1"]),
        ]
        for r in topic5_metrics
    ]
    add_table(doc, ["特征方案", "特征数", "样本数", "测试数", "Accuracy", "Macro F1"], perf_rows, "表4  选题5二分类性能对比")
    picture(doc, TOPIC5 / "figures" / "topic5_accuracy_macro_f1_bars.png", "图6  四种特征设置下的 Accuracy 与 Macro F1", 5.4)
    picture(doc, TOPIC5 / "figures" / "topic5_feature_space_comparison.png", "图7  四种特征空间中的美观/极致扭曲分布", 5.9)
    picture(doc, TOPIC5 / "figures" / "topic5_confusion_matrices.png", "图8  四种特征设置下的混淆矩阵", 5.6)

    top_rows = [[r["feature"], fmt(r["importance"], 4)] for r in top_features[:10]]
    add_table(doc, ["特征", "重要性"], top_rows, "表5  Fusion 模型前10个重要特征")
    picture(doc, TOPIC5 / "figures" / "topic5_fusion_top15_feature_importance.png", "图9  Fusion 模型 Top 15 特征重要性", 5.2)
    picture(doc, TOPIC5 / "figures" / "topic5_misclassified_examples.png", "图10  Fusion 模型典型误判样本", 5.8)
    paragraph(
        doc,
        "选题5结果显示，单独 ResNet50 与 Swin_T 特征均能捕捉一定整体风格差异，但准确率仍约在 0.68 至 0.70；Morphology 特征略高，说明可解释的轮廓与空间占用指标对扭曲识别更直接；Fusion 融合特征达到最高性能，说明深度视觉特征与形态学特征具有互补性。",
    )

    heading(doc, "5 AI 识别局限与实验边界", 1)
    paragraph(
        doc,
        "AI 模型在区分端正易读与极致扭曲这类差异较大的目标时表现较好，但在美观与极致扭曲之间仍存在局限。二者都可能包含曲线夸张、结构调整和个性化笔画，模型看到的是连续风格变化，而不是清晰类别边界。深度特征能判断整体视觉相似性，却不直接说明差异来自外接框扩张、轮廓复杂化、重心偏移还是参考符号偏离。因此，本项目引入 Morphology 和 Fusion 特征，以补充解释性。",
    )
    paragraph(
        doc,
        "实验仍存在边界：数据来自同一课程扫描模板，裁切参数较固定，模型可能学习到部分页面、网格或书写者习惯。虽然当前整理脚本已按 page_id 分组划分以降低同页泄漏风险，但后续仍可通过更多书写者、更多符号和跨人划分进一步检验泛化能力。",
    )

    heading(doc, "6 结论", 1)
    paragraph(
        doc,
        "本报告完成选题4并进一步拓展选题5。选题4证明端正易读、美观、极致扭曲三类书写目标在视觉特征空间中存在可测量差异；选题5进一步说明，美观与极致扭曲是更难的细粒度识别任务。深度视觉特征有助于捕捉整体风格，形态学特征有助于解释扭曲机制，融合特征则在性能和解释性之间取得更稳定的平衡。总体来看，AI 可以识别符号变形趋势，但对隐含审美边界和主动扭曲策略的理解仍有限。",
    )

    heading(doc, "参考文献", 1)
    refs = [
        "He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition[C]. CVPR, 2016.",
        "Liu Z, Lin Y, Cao Y, et al. Swin Transformer: Hierarchical vision transformer using shifted windows[C]. ICCV, 2021.",
        "Breiman L. Random forests[J]. Machine Learning, 2001.",
        "Kruskal W H, Wallis W A. Use of ranks in one-criterion variance analysis[J]. Journal of the American Statistical Association, 1952.",
    ]
    for i, ref in enumerate(refs, 1):
        paragraph(doc, f"[{i}] {ref}", indent=False, size=9.8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    if OUT.exists():
        backup = OUT.with_name(f"{OUT.stem}_旧版备份_{datetime.now():%Y%m%d_%H%M%S}{OUT.suffix}")
        try:
            shutil.copy2(OUT, backup)
            print(f"backup: {backup}")
        except PermissionError:
            print("backup skipped: target appears locked")
    try:
        doc.save(OUT)
        print(f"report: {OUT}")
    except PermissionError:
        doc.save(FALLBACK_OUT)
        print(f"report: {FALLBACK_OUT}")


if __name__ == "__main__":
    build()
