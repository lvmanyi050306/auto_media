# -*- coding: utf-8 -*-
"""Build a DOCX report that keeps the course template structure."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "智能媒体工作坊_技术报告模板..docx"
OUT = ROOT / "基础案例A-视觉分析2" / "技术报告_选题4_5_视觉模型拓展_严格模板版.docx"
TOPIC4 = ROOT / "outputs" / "topic4"
TOPIC5 = ROOT / "outputs" / "topic5"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"


def csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str | float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def set_run_font(run, size: float | None = None, bold: bool | None = None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def normalize_paragraph(p, size: float = 10.5):
    for run in p.runs:
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)


def set_text(p, text: str, size: float = 10.5, align=None):
    p.text = text
    if align is not None:
        p.alignment = align
    normalize_paragraph(p, size=size)
    return p


def insert_paragraph_after(element, doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    p.text = text
    if align is not None:
        p.alignment = align
    normalize_paragraph(p)
    element.addnext(p._p)
    return p


def insert_body_after(element, doc: Document, text: str):
    p = insert_paragraph_after(element, doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    return p._p


def insert_picture_after(element, doc: Document, image_path: Path, caption: str, width: float = 5.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        p.add_run().add_picture(str(image_path), width=Inches(width))
    else:
        p.add_run(f"（缺少图像：{image_path.name}）")
    element.addnext(p._p)
    normalize_paragraph(p)
    cap = insert_paragraph_after(p._p, doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in cap.runs:
        set_run_font(run, size=9)
    return cap._p


def insert_picture_only_after(element, doc: Document, image_path: Path, width: float = 5.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        p.add_run().add_picture(str(image_path), width=Inches(width))
    else:
        p.add_run(f"（缺少图像：{image_path.name}）")
    element.addnext(p._p)
    normalize_paragraph(p)
    return p._p


def insert_table_after(element, doc: Document, headers: list[str], rows: list[list[object]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    element.addnext(table._tbl)
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True)
        shade(cell, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cell, value, align=align)
    table_borders(table)
    return table._tbl


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def write_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_run_font(r, size=8.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_borders(table):
    pr = table._tbl.tblPr
    borders = pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C3D0")


def morphology_table_rows() -> list[list[str]]:
    return [
        ["ink_ratio", "墨迹像素占图像面积的比例", "反映书写浓度、笔画覆盖范围"],
        ["bbox_width_ratio", "墨迹外接框宽度占图像宽度的比例", "反映横向空间扩张或拉伸"],
        ["bbox_height_ratio", "墨迹外接框高度占图像高度的比例", "反映纵向拉伸或压缩"],
        ["bbox_area_ratio", "墨迹外接框面积占整图面积的比例", "反映整体空间占用扩大"],
        ["aspect_ratio", "外接框宽高比", "反映符号整体比例变化"],
        ["perimeter_norm", "归一化轮廓周长", "反映轮廓延展与边界复杂化"],
        ["complexity", "由轮廓周长与墨迹面积计算的复杂度", "反映笔画曲折、结构夸张程度"],
        ["edge_density", "边缘像素密度", "反映局部边界与细节变化"],
        ["centroid_x", "墨迹重心的横向位置", "反映结构左右偏移"],
        ["centroid_y", "墨迹重心的纵向位置", "反映结构上下偏移"],
        ["direction_entropy", "笔画方向分布的信息熵", "反映笔画方向变化和方向混杂程度"],
        ["ref_chamfer", "与参考符号轮廓的 Chamfer 距离", "反映相对参考符号的形态偏离"],
    ]


def fill_template(doc: Document):
    p = doc.paragraphs
    topic4_metrics = csv_rows(TOPIC4 / "metrics" / "topic4_classification_metrics.csv")[0]
    topic4_dist = csv_rows(TOPIC4 / "tables" / "topic4_group_distances.csv")
    topic5_perf = csv_rows(TOPIC5 / "tables" / "topic5_model_performance.csv")
    top_features = csv_rows(TOPIC5 / "tables" / "topic5_fusion_feature_importance.csv")[:10]
    fusion = next(row for row in topic5_perf if row["model"] == "Fusion")

    # Cover: do not change the user's intentional ④ + ⑤ checkbox choice.
    set_text(p[3], "《从端正到扭曲：不同书写目标下的符号变形规律与 AI 识别局限》", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_text(p[8], "请在以下选题中勾选一项（在 □ 内打 √）：", size=10.5)
    set_text(p[12], "☑ ④ 不同书写目标（端正/美观/扭曲）下的字迹差异", size=10.5)
    set_text(p[13], "☑ ⑤ 自定义拓展选题：从端正到扭曲：不同书写目标下的符号变形规律与 AI 识别局限（需先完成前4项之一）", size=10.5)

    set_text(
        p[16],
        f"本报告围绕不同书写目标下的符号变形规律与 AI 识别局限展开。研究首先完成选题4，基于端正易读、美观书写、极致扭曲三类扫描页，经裁切获得每类 2880 张单个抄写格图像，并复用基础案例A视觉分析流程，使用本地 ResNet50 权重提取深度视觉特征，进行 PCA、t-SNE 可视化和三分类检验。结果显示，三类书写目标在视觉特征空间中存在可测量差异，三分类 Accuracy 为 {fmt(topic4_metrics['accuracy'])}，Macro F1 为 {fmt(topic4_metrics['macro_f1'])}；其中美观与极致扭曲的组间距离最小。基于此，选题5进一步聚焦美观 vs 极致扭曲这一细粒度任务，并以 Morphology 成功提取特征的 5323 个有效样本为统一样本集合，比较 ResNet50、Swin_T、Morphology 与 Fusion 四种特征方案。结果显示 Fusion 模型 Accuracy 为 {fmt(fusion['accuracy'])}，Macro F1 为 {fmt(fusion['macro_f1'])}，优于单独深度特征和单独形态学特征。研究表明，深度视觉特征能捕捉整体视觉风格，形态学特征更有助于解释扭曲机制，二者融合可提升识别稳定性与可解释性；同时，AI 对“美观”和“极致扭曲”这种隐含审美边界仍存在识别局限。",
    )
    set_text(p[17], "关键词：符号变形；书写目标；ResNet50；Swin Transformer；形态学特征；AI识别局限")
    set_text(
        p[19],
        "This report studies symbol deformation patterns under different handwriting goals and the recognition limitations of AI visual models. Topic 4 verifies that legible, aesthetic, and extremely distorted writing styles show measurable differences in the visual feature space. Based on this result, Topic 5 focuses on the more difficult fine-grained task of distinguishing aesthetic writing from extremely distorted writing. Four feature settings are compared on the same valid sample set: ResNet50, Swin_T, morphology features, and fused features. The results indicate that deep features capture global visual style, while morphology features provide more interpretable cues about deformation mechanisms. Feature fusion achieves the most stable performance, but AI still has limitations in recognizing this implicit aesthetic boundary.",
    )
    set_text(p[20], "Keywords: symbol deformation; handwriting goals; ResNet50; Swin Transformer; morphology features; AI recognition limitations")

    set_text(p[23], "本研究关注同一符号在不同书写目标下产生的形态差异。端正易读强调结构稳定和可辨认性，美观书写强调比例、节奏和视觉张力，极致扭曲则主动改变常规形态但仍保留一定可识别线索。课程案例A已经提供了按 Data 子文件夹分组进行视觉特征分析的流程，因此本报告在该流程基础上完成选题4，并进一步把选题5作为选题4的升级拓展，分析美观与极致扭曲为何更难区分，以及 AI 视觉模型在理解隐含审美边界时的局限。")
    set_text(p[25], "相关理论主要包括 Polanyi 的隐性知识理论、笔迹形态分析和深度视觉特征分析。隐性知识强调许多判断难以完全由显性规则表达，美观与扭曲的边界也具有这种特征。深度视觉模型可以通过高维特征捕捉整体图像风格，但其输出不必然具有可解释性。因此，本项目同时引入外接框、墨迹占比、轮廓复杂度、方向熵和参考偏离等形态学指标，用于补充解释 AI 模型的分类依据。")
    set_text(p[28], "数据来自三类本地扫描页：端正易读方式抄写符号、美观方式抄写符号、极致扭曲方式抄写符号。实验没有直接使用整页扫描图，而是通过固定网格参数裁切为单个抄写格图像，并整理到基础案例A的 Data 目录下。当前 Data 结构包含端正易读、美观、极致扭曲三个子文件夹，每类 2880 张，共 8640 张。裁切索引保存为 Data/image_index.csv，字段包括 image_path、label、source_page、page_id、crop_id、row 和 col。")
    set_text(p[30], "选题4使用本地 ResNet50 权重提取深度视觉特征，并进行 PCA、t-SNE、组间距离和 RandomForest 三分类分析。选题5仅选择美观与极致扭曲两类，比较 ResNet50 深度特征、Swin_T 深度特征、Morphology 可解释形态学特征和 Fusion 融合特征。ResNet50 与 Swin_T 均优先加载本地模型权重；深度特征提取后通过 PCA 保留前 20 个主成分。t-SNE 仅用于二维可视化，不参与分类训练。分类器均使用 RandomForestClassifier，random_state 固定为 42，训练测试划分为 8:2，并使用 GroupShuffleSplit 按 page_id 划分，以降低同一扫描页样本同时进入训练集和测试集的风险。选题5四种特征方案均以 Morphology 成功提取特征的有效样本列表为基准，并共用同一训练集/测试集划分。")

    cap = insert_paragraph_after(p[30]._p, doc, "表1  形态学特征说明", align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_table_after(cap._p, doc, ["特征名称", "含义", "对应的字迹变化"], morphology_table_rows())

    set_text(p[32], "实验结果分为选题4和选题5两部分。选题4用于证明三类书写目标存在整体视觉差异；选题5在此基础上聚焦最容易重叠的美观与极致扭曲，进一步比较不同特征方案的识别能力和解释性。")
    set_text(p[33], "表2  数据整理结果")
    anchor = insert_table_after(
        p[33]._p,
        doc,
        ["类别", "裁切图像数", "原始扫描页来源", "整理后目录"],
        [
            ["端正易读", "2880", "3-以端正易读方式抄写符号", "Data/端正易读"],
            ["美观", "2880", "4-以美观方式抄写符号", "Data/美观"],
            ["极致扭曲", "2880", "5-以极致扭曲方式抄写符号", "Data/极致扭曲"],
        ],
    )
    set_text(p[34], "图1  三类书写目标裁切样例")
    insert_picture_only_after(anchor, doc, TOPIC4 / "figures" / "topic4_sample_examples.png", width=5.6)
    anchor = insert_picture_after(p[34]._p, doc, TOPIC4 / "figures" / "topic4_resnet50_pca_scatter.png", "图2  选题4 ResNet50 PCA 二维分布", width=5.4)
    anchor = insert_picture_after(anchor, doc, TOPIC4 / "figures" / "topic4_resnet50_tsne_scatter.png", "图3  选题4 ResNet50 t-SNE 二维分布", width=5.4)

    dist_rows = [[r["pair"], fmt(r["centroid_distance_top10PC"]), fmt(r["centroid_distance_tSNE"]), fmt(r["PC1_cliffs_delta"])] for r in topic4_dist]
    cap = insert_paragraph_after(anchor, doc, "表3  选题4三类书写目标视觉特征距离", align=WD_ALIGN_PARAGRAPH.CENTER)
    anchor = insert_table_after(cap._p, doc, ["对比", "前10 PC中心距离", "t-SNE中心距离", "PC1 Cliff's delta"], dist_rows)
    cap = insert_paragraph_after(anchor, doc, "表4  选题4三分类模型性能", align=WD_ALIGN_PARAGRAPH.CENTER)
    anchor = insert_table_after(
        cap._p,
        doc,
        ["模型", "样本数", "测试数", "Accuracy", "Macro F1"],
        [["ResNet50 + PCA + RandomForest", topic4_metrics["n_samples"], topic4_metrics["n_test"], fmt(topic4_metrics["accuracy"]), fmt(topic4_metrics["macro_f1"])]],
    )
    anchor = insert_picture_after(anchor, doc, TOPIC4 / "figures" / "topic4_confusion_matrix.png", "图4  选题4三分类混淆矩阵", width=4.4)

    cap = insert_paragraph_after(anchor, doc, "表5  选题5美观 vs 极致扭曲二分类性能", align=WD_ALIGN_PARAGRAPH.CENTER)
    perf_rows = [[r["model"], r["n_features"], r["n_samples"], r["n_test"], fmt(r["accuracy"]), fmt(r["macro_f1"])] for r in topic5_perf]
    anchor = insert_table_after(cap._p, doc, ["特征方案", "特征数", "样本数", "测试数", "Accuracy", "Macro F1"], perf_rows)
    anchor = insert_picture_after(anchor, doc, TOPIC5 / "figures" / "topic5_accuracy_macro_f1_bars.png", "图5  选题5四种特征设置下的 Accuracy 与 Macro F1", width=5.2)
    anchor = insert_picture_after(anchor, doc, TOPIC5 / "figures" / "topic5_feature_space_comparison.png", "图6  选题5四种特征空间分布对比", width=5.8)
    anchor = insert_picture_after(anchor, doc, TOPIC5 / "figures" / "topic5_confusion_matrices.png", "图7  选题5四种特征设置下的混淆矩阵", width=5.4)
    cap = insert_paragraph_after(anchor, doc, "表6  Fusion模型前10个重要特征", align=WD_ALIGN_PARAGRAPH.CENTER)
    anchor = insert_table_after(cap._p, doc, ["特征", "重要性"], [[r["feature"], fmt(r["importance"], 4)] for r in top_features])
    anchor = insert_picture_after(anchor, doc, TOPIC5 / "figures" / "topic5_fusion_top15_feature_importance.png", "图8  Fusion模型Top 15特征重要性", width=5.2)
    anchor = insert_picture_after(anchor, doc, TOPIC5 / "figures" / "topic5_misclassified_examples.png", "图9  Fusion模型典型误判样本", width=5.6)
    insert_body_after(
        anchor,
        doc,
        "典型误判样本分析：美观样本被误判为极致扭曲，通常与笔画拉伸明显、外接框扩大、结构夸张、轮廓复杂度较高以及与参考符号偏离较大有关；这些样本虽然仍可被人判断为具有美观意图，但在模型特征空间中更接近主动扭曲。极致扭曲样本被误判为美观，往往是因为其仍保留较完整的整体结构，局部扭曲程度不足，笔画节奏接近美观书写，或者轮廓虽有变化但仍具有较强可读性。由此可见，AI 并不是简单依据“像不像原符号”进行判断，而是在整体风格、空间占用、轮廓复杂度和参考偏离之间进行综合判断。这说明“美观”和“极致扭曲”之间存在连续风格变化，AI 对这种隐含审美边界仍存在识别局限。",
    )

    set_text(p[36], f"选题4结果说明，不同书写目标确实会在深度视觉特征空间中形成差异，其中端正易读与极致扭曲距离最大，美观与极致扭曲距离最小。选题5进一步表明，美观与极致扭曲并不是简单的好坏或规则/不规则关系，而是存在连续风格变化。在统一 5323 个有效样本并使用相同划分后，ResNet50、Swin_T、Morphology 和 Fusion 的 Accuracy 分别为 {fmt(topic5_perf[0]['accuracy'])}、{fmt(topic5_perf[1]['accuracy'])}、{fmt(topic5_perf[2]['accuracy'])} 和 {fmt(topic5_perf[3]['accuracy'])}。深度视觉特征能够捕捉整体视觉风格，但对外接框扩张、轮廓复杂化、重心偏移和参考符号偏离等具体机制解释不足。Morphology 特征提供了更直接的解释线索，而 Fusion 特征结合了整体风格和局部形态，因此表现更稳定。")
    set_text(p[38], "本报告严格按照课程模板完成选题4与选题5的合并呈现。选题4证明端正易读、美观、极致扭曲三类书写目标存在可测量视觉差异；选题5不是独立新题，而是选题4的升级拓展。相比选题4只回答“三类书写目标是否存在差异”，选题5进一步回答“差异来自哪里”以及“AI 为什么难以区分美观与极致扭曲”。结论是：AI 视觉模型可以识别符号变形趋势，但对隐含审美边界和主动扭曲策略的理解仍有限；将深度视觉特征与可解释形态学特征融合，可以提升识别稳定性和报告解释力。未来可加入更多书写者、更多符号类别以及跨人划分实验，以进一步验证泛化能力。")

    set_text(p[40], "[1] Polanyi M. The Tacit Dimension[M]. London: Routledge & Kegan Paul, 1966.")
    set_text(p[41], "[2] He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition[C]. CVPR, 2016.")
    set_text(p[42], "[3] Liu Z, Lin Y, Cao Y, et al. Swin Transformer: Hierarchical vision transformer using shifted windows[C]. ICCV, 2021.")
    insert_paragraph_after(p[42]._p, doc, "[4] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.")
    set_text(p[44], "补充材料包括：src/main.py、src/crop_pages.py、src/train_models.py、outputs/topic4、outputs/topic5、docs/报告补充材料.md。原始扫描页未直接进入模型分析，已裁切为单个抄写格图像；原始视觉分析.py 已备份为视觉分析_backup.py。")

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            for run in para.runs:
                set_run_font(run, bold=True)


def build():
    doc = Document(TEMPLATE)
    fill_template(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
