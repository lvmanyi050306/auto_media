# -*- coding: utf-8 -*-
"""Build the technical report DOCX from the analysis outputs."""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPORT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = REPORT_ROOT.parent
RESULTS_DIR = REPORT_ROOT / "results"
REPORT_DIR = REPORT_ROOT / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

TEMPLATE = next(WORKSPACE.glob("智能媒体工作坊_技术报告模板*.docx"))
OUT_DOCX = REPORT_DIR / "技术报告_选题5_基于题4拓展分析.docx"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING_COLOR = RGBColor(31, 78, 121)
ACCENT_FILL = "E8EEF5"


FEATURE_CN = {
    "ink_ratio": "墨迹占比",
    "bbox_area_ratio": "外接框面积比",
    "complexity": "轮廓复杂度",
    "elongation": "主轴细长度",
    "orientation_entropy": "方向熵",
    "ref_iou": "与参考重合度",
    "ref_chamfer": "参考偏离距离",
}

TABLE_FEATURES = [
    "ink_ratio",
    "bbox_area_ratio",
    "complexity",
    "orientation_entropy",
    "ref_iou",
    "ref_chamfer",
]

STYLE_ORDER = ["legible", "aesthetic", "distorted"]
STYLE_LABEL = {
    "legible": "端正易读",
    "aesthetic": "美观",
    "distorted": "极致扭曲",
}


def set_run_font(run, size=None, bold=None, color=None, cn_font=CN_FONT):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, size=10.5, align=None, bold=False, color=None):
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C3D0")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")


def add_heading(doc, text: str, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}" if level <= 3 else "Normal"
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 14, 2: 12, 3: 11}.get(level, 10.5), bold=True, color=HEADING_COLOR)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc, text: str, first_line_indent=True):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def add_caption(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=False, color=RGBColor(89, 89, 89))
    return paragraph


def add_picture(doc, path: Path, caption: str, width_inches: float):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def read_summary():
    summary = {}
    with (RESULTS_DIR / "summary_by_style.csv").open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            summary[(row["feature"], row["style"])] = row
    return summary


def read_tests():
    tests = {}
    with (RESULTS_DIR / "statistical_tests.csv").open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            tests[(row["feature"], row["pair"])] = row
    return tests


def read_counts():
    counts = {}
    with (RESULTS_DIR / "sample_counts.csv").open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            counts[row["style"]] = row
    return counts


def read_classification_metrics():
    text = (RESULTS_DIR / "classification_metrics.txt").read_text(encoding="utf-8")
    sections = {}
    current = None
    for line in text.splitlines():
        if "Three-style" in line:
            current = "three"
            sections[current] = {}
        elif "Aesthetic-vs-distorted" in line:
            current = "binary"
            sections[current] = {}
        elif current and line.startswith("Accuracy:"):
            sections[current]["accuracy"] = float(line.split(":", 1)[1].strip())
        elif current and line.startswith("Macro F1:"):
            sections[current]["macro_f1"] = float(line.split(":", 1)[1].strip())
    return sections


def format_p(value: str) -> str:
    p = float(value)
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def format_float(value: str | float, digits=3) -> str:
    return f"{float(value):.{digits}f}"


def fill_cover(doc: Document):
    paragraphs = doc.paragraphs
    paragraphs[3].text = "从美观到极致扭曲：基于题四的图形符号字迹变形拓展分析"
    paragraphs[8].text = "本报告选题如下："
    paragraphs[12].text = "□ ④ 不同书写目标（端正/美观/扭曲）下的字迹差异（作为基础分析完成）"
    paragraphs[13].text = "☑ ⑤ 自定义拓展选题：以选题④为基础，研究美观到极致扭曲的可量化变形机制"
    for idx in range(0, min(14, len(paragraphs))):
        p = paragraphs[idx]
        if idx <= 3:
            set_paragraph_format(p, size=16 if idx in (0, 2, 3) else 12, align=WD_ALIGN_PARAGRAPH.CENTER, bold=idx in (0, 2, 3))
        else:
            set_paragraph_format(p, size=11)


def remove_placeholder_body(doc: Document):
    for paragraph in list(doc.paragraphs)[14:]:
        delete_paragraph(paragraph)


def add_metric_table(doc: Document, summary, tests):
    add_caption(doc, "表1  关键形态指标均值与“美观 vs 极致扭曲”效应量")
    table = doc.add_table(rows=1, cols=7)
    style_table(table, [1.35, 0.82, 0.82, 0.95, 0.95, 0.70, 0.65])
    headers = ["指标", "端正", "美观", "极致扭曲", "Cliff's δ", "p", "方向"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, ACCENT_FILL)
    for feature in TABLE_FEATURES:
        row = table.add_row()
        values = [
            FEATURE_CN[feature],
            format_float(summary[(feature, "legible")]["mean"]),
            format_float(summary[(feature, "aesthetic")]["mean"]),
            format_float(summary[(feature, "distorted")]["mean"]),
            format_float(tests[(feature, "美观 vs 极致扭曲")]["cliffs_delta"]),
            format_p(tests[(feature, "美观 vs 极致扭曲")]["pair_p"]),
            tests[(feature, "美观 vs 极致扭曲")]["direction"],
        ]
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, value, size=9, align=align)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("注：δ为“美观相对极致扭曲”的Cliff's delta；负值表示极致扭曲组更高。")
    set_run_font(run, size=8.5, color=RGBColor(89, 89, 89))


def add_classification_table(doc: Document, metrics):
    add_caption(doc, "表2  基于形态特征的书写目标识别结果")
    table = doc.add_table(rows=1, cols=4)
    style_table(table, [2.0, 1.0, 1.0, 2.4])
    headers = ["任务", "Accuracy", "Macro F1", "解释"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, ACCENT_FILL)
    rows = [
        ("三分类：端正/美观/极致扭曲", metrics["three"]["accuracy"], metrics["three"]["macro_f1"], "极致扭曲较容易识别；端正与美观重叠较多。"),
        ("二分类：美观/极致扭曲", metrics["binary"]["accuracy"], metrics["binary"]["macro_f1"], "拓展题核心任务，可达到约72%的可分辨度。"),
    ]
    for task, acc, f1, interp in rows:
        row = table.add_row()
        vals = [task, f"{acc:.3f}", f"{f1:.3f}", interp]
        for idx, (cell, value) in enumerate(zip(row.cells, vals)):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, value, size=9, align=align)


def build_report():
    summary = read_summary()
    tests = read_tests()
    counts = read_counts()
    metrics = read_classification_metrics()

    doc = Document(TEMPLATE)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    fill_cover(doc)
    remove_placeholder_body(doc)
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("摘  要")
    set_run_font(run, size=14, bold=True, color=HEADING_COLOR)
    add_body(
        doc,
        "本研究以选题④“三种书写目标下的字迹差异”为基础，进一步提出选题⑤的拓展问题：当书写者从“美观”转向“极致扭曲”时，哪些可量化形态维度最能解释这种变形，且AI或统计模型能否稳定识别这种书写意图？研究从端正易读、美观、极致扭曲三套扫描页中裁切抄写格，得到8640个样本；剔除近空白格后保留8057个有效样本。通过墨迹占比、外接框面积、轮廓复杂度、方向熵、参考偏离距离等特征，结合Kruskal-Wallis检验、Mann-Whitney检验、Cliff's delta效应量与随机森林分类实验进行分析。结果显示，极致扭曲组在外接框面积、复杂度、方向熵和参考偏离距离上显著高于美观组；美观与端正易读之间则高度重叠。三分类准确率约为0.540，而美观与极致扭曲二分类准确率约为0.718，说明极端变形意图具有可识别的形态信号，但普通“好看”与“端正”的边界更多依赖隐性审美判断。",
    )
    add_body(doc, "关键词：字迹变形；图形符号；隐性知识；形态特征；随机森林", first_line_indent=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    set_run_font(run, size=14, bold=True, color=HEADING_COLOR)
    add_body(
        doc,
        "This report extends Topic 4 by asking which measurable deformation dimensions separate aesthetic copying from extreme distortion. A fixed-grid cropper extracted copied-symbol cells from three scanned datasets. After filtering nearly blank cells, 8,057 valid samples were analyzed with morphology features, reference-distance measures, non-parametric tests, effect sizes, PCA visualization, and random-forest classifiers. Extreme distortion showed larger spatial spread, higher contour complexity, higher orientation entropy, and larger distance from the reference symbol. Three-way classification reached 0.540 accuracy, while aesthetic-vs-distorted classification reached 0.718. The findings suggest that extreme distortion produces detectable visual signals, whereas the boundary between legibility and beauty remains more tacit and ambiguous.",
        first_line_indent=False,
    )
    add_body(doc, "Keywords: handwriting deformation; graphic symbols; tacit knowledge; morphology; random forest", first_line_indent=False)

    add_heading(doc, "1 引言", 1)
    add_body(
        doc,
        "选题④要求比较端正易读、美观、极致扭曲三种方式下的字迹变形差异。若只停留在“是否显著不同”，结论容易变成统计显著而解释不足。因此本报告将选题⑤定义为选题④的升级：不仅比较三组是否不同，还追问差异主要来自哪些可解释的视觉维度，以及这种差异是否足以让模型反推出书写目标。",
    )
    add_body(
        doc,
        "从隐性知识角度看，书写者知道如何把一个符号写得“好看”或“扭曲”，但这类判断未必能被完整说成规则。将扫描页转化为形态指标，等于把部分隐性书写经验外化为可测量的空间、轮廓与方向特征；而分类实验则检验这些外化指标能否支撑机器识别。",
    )

    add_heading(doc, "2 数据来源与预处理", 1)
    add_body(
        doc,
        "数据来自三套扫描页：端正易读、美观、极致扭曲。每套包含40张扫描页，页面尺寸均为1536×2184像素。页面采用左右两栏、每栏12行的固定版式；每行包含参考符号和三次抄写格。本研究只分析三次抄写格，避免题号、页眉和参考列直接参与统计。",
    )
    add_body(
        doc,
        f"裁切后理论样本量为8640个。考虑到任务允许“写2~3遍”，少量第三格为空或墨迹极浅，本研究以墨迹占比不低于{0.005:.3f}作为有效样本阈值。有效样本数分别为：端正易读{counts['legible']['valid_samples']}个，美观{counts['aesthetic']['valid_samples']}个，极致扭曲{counts['distorted']['valid_samples']}个，共8057个。",
    )
    add_picture(
        doc,
        RESULTS_DIR / "example_crops_contact_sheet.png",
        "图1  三种书写目标下的参考格与抄写格示例（红框为参考，绿框为抄写）",
        5.35,
    )

    add_heading(doc, "3 方法", 1)
    add_heading(doc, "3.1 形态特征", 2)
    add_body(
        doc,
        "每个抄写格先转为灰度图，使用Otsu阈值并结合小连通域过滤得到墨迹掩膜。随后提取两类特征：其一是单格内部形态，包括墨迹占比、外接框面积、轮廓复杂度、主轴细长度、方向熵、连通域数和边缘密度；其二是与同一行参考符号的相对偏离，包括重合度、余弦相似度和对称Chamfer距离。这样既能衡量符号本身写得多大、多散、多复杂，也能衡量它偏离原符号的程度。",
    )
    add_heading(doc, "3.2 统计与识别实验", 2)
    add_body(
        doc,
        "由于各指标分布并非严格正态，本研究使用Kruskal-Wallis检验判断三组总体差异，并使用Mann-Whitney U检验与Cliff's delta比较美观与极致扭曲之间的效应方向和大小。为了检验差异的可识别性，进一步用随机森林进行两类任务：三分类识别端正/美观/极致扭曲，以及二分类识别美观/极致扭曲。训练测试按75%/25%分层划分，随机种子固定为42。",
    )

    add_heading(doc, "4 结果", 1)
    add_picture(
        doc,
        RESULTS_DIR / "feature_boxplots.png",
        "图2  三种书写目标下关键形态特征的分布",
        6.35,
    )
    add_body(
        doc,
        "图2显示，极致扭曲组在墨迹占比、外接框面积、轮廓复杂度和方向熵上整体上移，说明扭曲并非单纯“写歪”，而是同时表现为空间展开、轮廓增殖和方向变化更杂。美观组与端正易读组的分布更接近，说明“美观”更多是在保持结构稳定的基础上调整笔画粗细、曲率和姿态。",
    )
    add_metric_table(doc, summary, tests)
    add_body(
        doc,
        "表1进一步表明，美观与极致扭曲在所有列出的关键指标上均达到显著差异。最强的效应来自外接框面积比（δ=-0.301），其次是轮廓复杂度、方向熵和参考偏离距离。与之相反，主轴细长度和参考重合度在美观组略高，说明美观书写往往更倾向于保持清晰的主方向和与参考符号的结构一致性。",
    )
    add_picture(
        doc,
        RESULTS_DIR / "pca_scatter_features.png",
        "图3  形态特征的PCA二维投影",
        5.75,
    )
    add_body(
        doc,
        "PCA投影中三类样本并没有形成完全分离的簇。极致扭曲样本在边缘区域更常见，但端正与美观大面积交叠。这说明三种书写目标并不是离散风格标签，而是在连续形态空间中发生偏移；人的审美判断能利用局部结构、语境和经验，而简单形态指标只能捕捉其中一部分。",
    )
    add_classification_table(doc, metrics)
    add_picture(
        doc,
        RESULTS_DIR / "feature_importance_aesthetic_vs_distorted.png",
        "图4  美观与极致扭曲二分类中的前10个重要特征",
        5.85,
    )
    add_body(
        doc,
        "分类结果支持上述解释。三分类准确率约0.540，只高于随机水平但并不充分，主要困难来自端正与美观的边界模糊；而聚焦题五的美观/极致扭曲二分类可达0.718，说明极致扭曲确实会留下较稳定的可量化信号。特征重要性显示，外接框宽度、轮廓复杂度、参考偏离距离和方向熵是区分二者的主要依据。",
    )

    add_heading(doc, "5 讨论", 1)
    add_body(
        doc,
        "题四的基础结论是：三种书写目标之间存在显著字迹差异，但差异大小不均衡。题五的拓展结论是：差异并不平均分布在所有指标上，而集中体现在“空间扩张—轮廓复杂化—方向多样化—参考偏离”这一链条。极致扭曲要求书写者保留关键结构特征，同时主动破坏常规形态，因此更容易造成外接框扩大、局部绕曲增多和方向熵升高。",
    )
    add_body(
        doc,
        "美观书写则更像一种约束优化：它并不追求最大变化，而是在可识别结构内调整笔画粗细、平衡、曲线和节奏。这也解释了为什么美观与端正易读在统计空间中高度重叠。换言之，极致扭曲是“显性策略”更强的目标，而美观是“隐性判断”更强的目标；后者需要人的经验、审美和上下文来判断，单纯几何指标很难完全覆盖。",
    )
    add_body(
        doc,
        "本研究仍有局限。第一，裁切采用固定网格，虽然扫描尺寸一致，但个别超出格子的笔画可能被截断。第二，研究未手工标注符号类别，因此没有进一步分析不同符号结构对变形策略的影响。第三，分类实验使用的是手工形态特征，而非深度视觉嵌入；若后续结合CLIP、ViT或符号结构标签，可能更好地解释哪些符号更容易在扭曲中保持可识别性。",
    )

    add_heading(doc, "6 结论与展望", 1)
    add_body(
        doc,
        "以选题④为基础，本报告完成了选题⑤的拓展研究。结论可概括为三点：第一，端正、美观、极致扭曲三种目标下的字迹变形存在显著差异；第二，美观到极致扭曲的主要变化不是单一指标上升，而是空间展开、复杂度增加、方向更分散和相对参考符号偏离共同发生；第三，模型能够较好地区分美观与极致扭曲，但难以稳定区分端正与美观，说明“美观”具有更强的隐性审美成分。",
    )
    add_body(
        doc,
        "后续可以沿两条路线继续推进：一是加入符号类别标签，分析复杂符号、闭合符号、线性符号在扭曲时的差异；二是引入深度视觉特征与人工评分，比较机器可测特征、人类审美判断和AI理解能力之间的偏差。",
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "Polanyi M. The Tacit Dimension[M]. London: Routledge & Kegan Paul, 1966.",
        "Otsu N. A threshold selection method from gray-level histograms[J]. IEEE Transactions on Systems, Man, and Cybernetics, 1979, 9(1): 62-66.",
        "Kruskal W H, Wallis W A. Use of ranks in one-criterion variance analysis[J]. Journal of the American Statistical Association, 1952, 47(260): 583-621.",
        "Mann H B, Whitney D R. On a test of whether one of two random variables is stochastically larger than the other[J]. Annals of Mathematical Statistics, 1947, 18(1): 50-60.",
        "Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_body(doc, f"[{idx}] {ref}", first_line_indent=False)

    # Compact the document a little for the course page budget.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
