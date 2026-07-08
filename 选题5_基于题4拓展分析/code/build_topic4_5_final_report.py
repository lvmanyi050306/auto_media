# -*- coding: utf-8 -*-
"""Build final combined report for Topic 4 and Topic 5."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
TOPIC4 = WORKSPACE / "基础案例A-视觉分析2"
TEMPLATE = next(WORKSPACE.glob("智能媒体工作坊_技术报告模板*.docx"))
OUT_DOCX = ROOT / "report" / "技术报告_选题4_5_视觉模型拓展.docx"
FALLBACK_DOCX = TOPIC4 / "技术报告_选题4_5_视觉模型拓展.docx"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "E8EEF5"


def run_font(run, size=10.5, bold=False, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def paragraph_text(doc, text, indent=True, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    run_font(r, size=size)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    run_font(r, size={1: 14, 2: 12}.get(level, 11), bold=True, color=HEADING)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    run_font(r, size=9, color=MUTED)
    return p


def add_picture(doc, path: Path, cap: str, width: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def cell_text(cell, text, bold=False, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_borders(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    pr = table._tbl.tblPr
    borders = pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C3D0")


def read_csv_dicts(path: Path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_topic4_metrics():
    pair = read_csv_dicts(TOPIC4 / "4_选题4统计检验" / "resnet50_pairwise_group_distances.csv")
    metrics_text = (TOPIC4 / "4_选题4统计检验" / "resnet50_classification_metrics.txt").read_text(encoding="utf-8")
    acc = "0.7407"
    f1 = "0.7372"
    for line in metrics_text.splitlines():
        if line.startswith("Accuracy:"):
            acc = line.split(":", 1)[1].strip()
        if line.startswith("Macro F1:"):
            f1 = line.split(":", 1)[1].strip()
    return pair, acc, f1


def read_topic5_metrics():
    return read_csv_dicts(TOPIC4 / "4_选题4统计检验" / "topic5_model_comparison_metrics.csv")


def fill_cover(doc):
    p = doc.paragraphs
    p[3].text = "不同书写目标下的字迹差异及多模型扭曲识别拓展分析"
    p[8].text = "本报告合并完成选题④，并在此基础上进一步完成选题⑤："
    p[12].text = "☑ ④ 不同书写目标（端正/美观/扭曲）下的字迹差异"
    p[13].text = "☑ ⑤ 自定义拓展选题：基于多模型视觉特征分析美观到极致扭曲的可分性与AI局限"
    for idx in range(14):
        para = p[idx]
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run_font(run, size=16 if idx in (0, 2, 3) else 11, bold=idx in (0, 2, 3))


def add_topic4_table(doc, pair_rows):
    caption(doc, "表1  选题4三组书写目标的视觉特征距离")
    table = doc.add_table(rows=1, cols=4)
    table_borders(table)
    headers = ["对比", "前10 PC中心距离", "t-SNE中心距离", "PC1 Cliff's δ"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell_text(cell, text, bold=True)
        shade(cell, TABLE_FILL)
    for row in pair_rows:
        cells = table.add_row().cells
        vals = [
            row["pair"],
            f"{float(row['centroid_distance_top10PC']):.3f}",
            f"{float(row['centroid_distance_tSNE']):.3f}",
            f"{float(row['PC1_cliffs_delta']):.3f}",
        ]
        for i, (cell, val) in enumerate(zip(cells, vals)):
            cell_text(cell, val, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)


def add_topic5_table(doc, metrics):
    caption(doc, "表2  选题5美观/极致扭曲二分类模型比较")
    table = doc.add_table(rows=1, cols=5)
    table_borders(table)
    headers = ["特征/模型", "样本数", "特征数", "Accuracy", "Macro F1"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell_text(cell, text, bold=True)
        shade(cell, TABLE_FILL)
    for row in metrics:
        cells = table.add_row().cells
        vals = [
            row["model"],
            row["n_samples"],
            row["n_features"],
            f"{float(row['accuracy']):.3f}",
            f"{float(row['macro_f1']):.3f}",
        ]
        for cell, val in zip(cells, vals):
            cell_text(cell, val)


def add_model_resource_table(doc):
    caption(doc, "表3  本地模型资源与选题5使用情况")
    table = doc.add_table(rows=1, cols=4)
    table_borders(table)
    headers = ["模型", "位置", "本次处理", "说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell_text(cell, text, bold=True, size=8.8)
        shade(cell, TABLE_FILL)
    rows = [
        ("ResNet50", "视觉特征计算/model/torchvision", "已使用", "CPU速度较快，用于题4和题5基线深度特征。"),
        ("Swin_T", "视觉特征计算/model/torchvision", "已使用", "Transformer视觉模型，用于题5二分类比较。"),
        ("ViT_B16", "视觉特征计算/model/torchvision", "未作为主实验", "CPU推理成本较高，适合后续补充。"),
        ("CLIP / DINOv2", "视觉特征计算/model/hf_models", "未作为主实验", "适合语义/自监督视觉特征扩展，但当前环境运行成本更高。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT if i in (1, 3) else WD_ALIGN_PARAGRAPH.CENTER)


def build():
    pair_rows, topic4_acc, topic4_f1 = read_topic4_metrics()
    topic5_metrics = read_topic5_metrics()

    doc = Document(TEMPLATE)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
    fill_cover(doc)
    for para in list(doc.paragraphs)[14:]:
        delete_paragraph(para)

    doc.add_page_break()
    heading(doc, "摘  要", 1)
    paragraph_text(
        doc,
        "本报告先完成选题④：基于三种书写目标（端正易读、美观、极致扭曲）的扫描页，裁切得到8640个抄写格，并沿用基础案例A的视觉分析流程，将图像按Data子文件夹分组，提取ResNet50深度视觉特征，进行PCA、t-SNE可视化和分类检验。结果显示三组在深度视觉空间中存在显著差异，三分类准确率为0.7407，端正易读与极致扭曲差异最大，美观与极致扭曲重叠最多。在此基础上，选题⑤进一步聚焦最难区分的“美观 vs 极致扭曲”，比较ResNet50、Swin_T、可解释形态特征及融合特征的识别能力。实验发现，单独深度特征准确率约0.67，形态特征约0.704，而融合模型达到0.739，说明AI视觉模型能捕捉整体视觉风格，但笔画外接框、复杂度、参考偏离距离等可解释特征对理解扭曲机制仍有重要补充。",
    )
    paragraph_text(doc, "关键词：字迹变形；视觉特征；ResNet50；Swin Transformer；隐性知识；模型融合", indent=False)

    heading(doc, "1 引言", 1)
    paragraph_text(doc, "选题④关注不同书写目标下字迹是否存在显著差异。选题⑤则在完成选题④之后继续追问：当端正易读、美观、极致扭曲已经被证明可分时，模型到底依据哪些视觉线索识别差异？尤其是“美观”和“极致扭曲”都可能包含笔画夸张、结构调整和形态变化，它们之间的边界比“端正 vs 扭曲”更模糊，因此更适合作为拓展研究对象。")

    heading(doc, "2 数据与预处理", 1)
    paragraph_text(doc, "原始数据包括端正易读、美观、极致扭曲三套扫描页，每套40页。每页左右两栏，每栏约12行，每行包含参考符号和三次抄写。为避免页眉、题号、网格和整页扫描噪声干扰，本报告使用固定网格裁切脚本提取抄写格，每组三种书写目标各2880张，共8640张，并放入基础案例A的Data目录下作为三个分组。")
    paragraph_text(doc, "选题4使用全部三组图像；选题5聚焦美观与极致扭曲两组，共5760张深度特征样本。形态学分析中剔除了近空白格，因此有效样本为5323张。")

    heading(doc, "3 选题4：三种书写目标的视觉差异", 1)
    paragraph_text(doc, "按照基础案例A的流程，Data目录下的每个子文件夹被视为一个组。由于原始视觉分析.py在当前环境中使用的旧版ViT/CLIP接口运行成本较高，本报告保留其分组、特征、PCA和t-SNE分析框架，使用本地可用的ResNet50权重完成深度视觉特征提取。")
    add_picture(doc, TOPIC4 / "3_降维到二维平面" / "resnet50_t-SNE_可视化.png", "图1  选题4三种书写目标的ResNet50特征t-SNE分布", 5.8)
    add_topic4_table(doc, pair_rows)
    paragraph_text(doc, f"ResNet50 PCA特征的三分类准确率为{float(topic4_acc):.3f}，Macro F1为{float(topic4_f1):.3f}。表1显示，端正易读与极致扭曲的中心距离最大，端正易读与美观次之，美观与极致扭曲最小。这说明三种目标存在系统差异，但美观与极致扭曲在视觉空间中并非完全分离。")

    heading(doc, "4 选题5拓展方向与模型选择", 1)
    paragraph_text(doc, "选题5的拓展方向确定为：基于多模型视觉特征分析“美观 vs 极致扭曲”的可分性与AI视觉局限。这个方向建立在选题4结论之上，专门研究最容易混淆的一组对比，并将深度视觉模型与可解释形态指标结合。")
    add_model_resource_table(doc)
    paragraph_text(doc, "本次实际使用ResNet50和Swin_T两个本地模型。ResNet50作为题4延续基线，Swin_T作为更强的分层Transformer视觉模型。与此同时，引入形态学特征，包括墨迹占比、外接框面积、轮廓复杂度、方向熵和参考偏离距离等，用于解释模型无法直接说明的笔画变形机制。")

    heading(doc, "5 选题5实验结果", 1)
    add_picture(doc, TOPIC4 / "4_选题4统计检验" / "topic5_model_accuracy_comparison.png", "图2  美观与极致扭曲二分类的模型识别效果", 5.8)
    add_topic5_table(doc, topic5_metrics)
    paragraph_text(doc, "从表2可见，ResNet50和Swin_T单独使用时准确率均约为0.67，说明深度模型可以捕捉一部分视觉风格差异，但两类仍大量重叠。形态学特征达到0.704，略高于单独深度模型，说明外接框、复杂度和参考偏离这类可解释指标对扭曲识别更直接。融合模型达到0.739，表明深度特征与形态特征具有互补性。")
    add_picture(doc, TOPIC4 / "4_选题4统计检验" / "topic5_feature_space_comparison.png", "图3  不同特征空间中的美观/极致扭曲分布", 6.1)
    add_picture(doc, TOPIC4 / "4_选题4统计检验" / "topic5_confusion_matrices.png", "图4  四种特征设置下的混淆矩阵", 6.1)
    add_picture(doc, TOPIC4 / "4_选题4统计检验" / "topic5_fusion_feature_importance.png", "图5  融合模型的重要特征", 5.8)
    paragraph_text(doc, "融合模型的重要特征显示，外接框宽度、外接框面积、轮廓复杂度、周长、参考Chamfer距离与Swin/ResNet的若干主成分共同贡献最大。这说明极致扭曲并不只是“写得不像”，而是表现为空间范围扩大、轮廓复杂化、局部方向变化增加以及相对参考符号的偏离。")

    heading(doc, "6 讨论", 1)
    paragraph_text(doc, "选题4证明了不同书写目标在视觉特征空间中的差异；选题5进一步显示，差异的可识别性取决于特征类型。深度视觉模型擅长捕捉整体图像风格，但在美观与极致扭曲这种相邻目标之间，单独深度特征不一定优于人工设计的形态指标。原因可能是两组都包含手写笔触、曲线和个性化变形，模型看到的是连续风格变化，而不是清晰类别边界。")
    paragraph_text(doc, "从隐性知识角度看，“美观”并非单纯的几何规则，而是书写者对平衡、节奏、粗细和可读性的综合判断；“极致扭曲”则是在保留关键结构的同时主动破坏常规形态。深度模型能发现二者有差异，但需要结合可解释指标才能说明差异来自哪里。")

    heading(doc, "7 结论", 1)
    paragraph_text(doc, "本报告完成了选题④并进一步完成选题⑤。选题④结果表明，端正易读、美观、极致扭曲三种书写目标存在显著视觉差异，且端正易读与极致扭曲差异最大。选题⑤结果表明，美观与极致扭曲是更具挑战的细粒度区分任务，单独ResNet50和Swin_T准确率约0.67，形态特征约0.704，融合模型约0.739。最终结论是：AI视觉模型能够识别字迹变形的整体趋势，但对“美观”和“扭曲”这种隐性审美边界仍存在局限；将深度特征与可解释形态特征结合，能得到更稳定也更可解释的分析。")

    heading(doc, "参考文献", 1)
    refs = [
        "Polanyi M. The Tacit Dimension[M]. London: Routledge & Kegan Paul, 1966.",
        "He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition[C]. CVPR, 2016.",
        "Liu Z, Lin Y, Cao Y, et al. Swin Transformer: Hierarchical vision transformer using shifted windows[C]. ICCV, 2021.",
        "Kruskal W H, Wallis W A. Use of ranks in one-criterion variance analysis[J]. Journal of the American Statistical Association, 1952.",
        "Breiman L. Random forests[J]. Machine Learning, 2001.",
    ]
    for i, ref in enumerate(refs, 1):
        paragraph_text(doc, f"[{i}] {ref}", indent=False, size=9.8)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_DOCX)
        print(OUT_DOCX)
    except PermissionError:
        doc.save(FALLBACK_DOCX)
        print(FALLBACK_DOCX)


if __name__ == "__main__":
    build()
